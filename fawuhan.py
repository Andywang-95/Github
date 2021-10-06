# 案例 1 ：批量生成法务函
# 目标：把【封号名单.xlsx】工作簿中的每个封号人员的名字和微信号添加到【法务函模板.docx】Word 文件对应的位置上，并将【法务函模板.docx】Word 文件另存为【法务函-XXX.docx】 Word 文件
# 请随时查看知识库和案例练习助手，与自己编写代码的步骤和内容比对参考，训练思维

# 方法一
# from openpyxl import load_workbook
# from docx import Document
# from docx.shared import Pt


# # 打开封号名单.xlsx
# wb = load_workbook('工作\封号名单.xlsx')
# ws = wb.active
# dict_name = {}

# # 将excel中的名字和微信号保存成键值对存入字典dict_name中
# for row in ws.iter_rows(min_row=2,values_only=True):
#     dict_name[row[0]] = row[1]
# print(len(dict_name))
# print(dict_name)

# # 遍历字典，取出名字和对应微信号
# for name in dict_name:
#     doc = Document('工作\法务函文件\法务函模板.docx')
#     para = doc.paragraphs[5]
#     # 添加名字并设置参数
#     run_name = para.add_run(name)
#     run_name.font.size = Pt(14)
#     run_name.font.bold = True
#     run_name.font.underline = True
#     # 添加微信号并设置参数
#     run_id = para.add_run(' 同学（WeChat ID：{}）'.format(dict_name[name]))
#     run_id.font.size = Pt(14)
#     # 保存文件
#     doc.save('工作\法务函文件\法务函-{}.docx'.format(name))
    



# 方法二
# from openpyxl import load_workbook
# from docx import Document
# from docx.shared import Pt


# # 打开封号名单.xlsx
# wb = load_workbook('工作\封号名单.xlsx')
# ws = wb.active

# # 将excel中的名字和微信号保存成键值对存入字典dict_name中
# for row in ws.iter_rows(min_row=2,values_only=True):
    
#     # 读取模版
#     doc = Document('工作\法务函文件\法务函模板.docx')
#     #读取对应段落
#     para = doc.paragraphs[5]

#     # 添加内文姓名并更改字号、以及相关数据
#     run_name = para.add_run(row[0])
#     run_name.font.size = Pt(14)
#     run_name.font.bold = True
#     run_name.font.underline = True
#     # 增加姓名後的内容，带入微信号
#     run_id = para.add_run(' 同学（WeChat ID：{}）'.format(row[1]))
#     run_id.font.size = Pt(14)
#     # 储存档案，因为姓名有重复的，所以在保存的时候档案名不能只添加姓名，否则会被覆盖掉，所以后面再添加了微信号以做区隔
#     doc.save('工作\法务函文件\法务函-{}-{}.docx'.format(row[0],row[1]))



# 方法三
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

# 定义添加所有元素的函数
def add_imfo(name,wechat):
    # 读取模版
    doc = Document('工作\法务函文件\法务函模板.docx')
    #读取对应段落
    para = doc.paragraphs[5]

    # 添加内文姓名并更改字号、以及相关数据
    run_name = para.add_run(name)
    run_name.font.size = Pt(14)
    run_name.font.bold = True
    run_name.font.underline = True
    # 增加姓名後的内容，带入微信号
    run_id = para.add_run(' 同学（WeChat ID：{}）'.format(wechat))
    run_id.font.size = Pt(14)
    # 储存档案，因为姓名有重复的，所以在保存的时候档案名不能只添加姓名，否则会被覆盖掉，所以后面再添加了微信号以做区隔
    doc.save('工作\法务函文件\法务函-{}-{}.docx'.format(name,wechat))

# 打开封号名单.xlsx
wb = load_workbook('工作\封号名单.xlsx')
ws = wb.active


# 定义主体，遍历excel内容，提出名字和微信号
def main():
    for row in ws.iter_rows(min_row=2,values_only=True):
        add_imfo(row[0],row[1])
    
# 呼叫主体
main()